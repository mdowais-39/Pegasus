"""Render the report data model to a .docx document.

Minimal, statistics-first and investigator-friendly. No charts — just the
numbers (in the summary) plus the two explained sections (round trips and money
flow). Every account identifier is already resolved to an account number / file
name upstream in report_builder, so no opaque statement UUID appears here.
"""

import io
from docx import Document


def _money(v):
    try:
        return f"Rs {float(v):,.0f}"
    except (TypeError, ValueError):
        return "-" if v in (None, "") else str(v)


def _heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _stat_table(doc, pairs):
    """Two-column Metric | Value table."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in pairs:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = "" if v is None else str(v)


def _table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = "" if val is None else str(val)


# Curated, plain-language metrics only (no jargon like "communities").
KEY_METRICS = [
    ("statements", "Statements"),
    ("transactions", "Transactions"),
    ("total_credit", "Total Credit"),
    ("total_debit", "Total Debit"),
    ("duplicates", "Duplicate Transactions"),
    ("failed_or_reversed", "Failed / Reversed"),
    ("round_trips_detected", "Round Trips Detected"),
    ("high_risk_accounts", "High-Risk Accounts"),
]
_MONEY_METRICS = {"total_credit", "total_debit"}

CATEGORY_LABELS = [
    ("total_transactions", "Total Transactions"), ("credits", "Credits"),
    ("debits", "Debits"), ("atm_withdrawals", "ATM Withdrawals"),
    ("cash_deposits", "Cash Deposits"), ("failed_transactions", "Failed / Reversed"),
    ("digital_upi", "Digital (UPI / apps)"), ("cheque", "Cheque"),
    ("card_pos", "Card / POS"),
]


def build_docx(report: dict) -> bytes:
    doc = Document()
    doc.add_heading(report.get("title", "Investigation Report"), level=0)
    doc.add_paragraph(report.get("scope") or f"Case: {report.get('case_id', 'all')}")
    if report.get("generated_at"):
        doc.add_paragraph(f"Generated: {report.get('generated_at')}")

    # ---------------------------------------------- Summary — key statistics
    _heading(doc, "Summary — Key Statistics")

    es = report.get("executive_summary", {}) or {}
    _stat_table(doc, [
        (lbl, _money(es.get(key)) if key in _MONEY_METRICS else es.get(key, 0))
        for key, lbl in KEY_METRICS
    ])

    dist = report.get("risk_distribution") or {}
    if any(dist.values()):
        _heading(doc, "Accounts by Risk Level", level=2)
        _stat_table(doc, [
            ("Critical", dist.get("CRITICAL", 0)), ("High", dist.get("HIGH", 0)),
            ("Medium", dist.get("MEDIUM", 0)), ("Low", dist.get("LOW", 0)),
        ])

    cats = report.get("category_counts") or {}
    if cats:
        _heading(doc, "Transaction Categories", level=2)
        _stat_table(doc, [(lbl, cats.get(key, 0)) for key, lbl in CATEGORY_LABELS])

    channels = report.get("channel_breakdown") or []
    if channels:
        _heading(doc, "Payment Channels", level=2)
        _table(doc, ["Channel / Class", "Transactions", "Total Value", "Share"],
               [[c.get("channel"), c.get("count"), _money(c.get("value")),
                 f"{round((c.get('share') or 0) * 100)}%"] for c in channels])

    # flagged findings live INSIDE the summary
    _heading(doc, "Flagged Findings", level=2)
    doc.add_paragraph(report.get("flags_summary", ""))
    flagged = report.get("flagged_findings", [])
    if flagged:
        _table(doc, ["Account", "Severity", "Flags", "Evidence"],
               [[f.get("account"), f.get("severity"),
                 ", ".join(f.get("tags", []) or []),
                 "; ".join(f.get("reasons", []) or [])] for f in flagged])

    # ------------------------------------------------------------ Round trips
    _heading(doc, "Round Trips (Circular Flows)")
    if report.get("round_trips_definition"):
        doc.add_paragraph().add_run(report["round_trips_definition"]).italic = True
    trips = report.get("round_trips", [])
    if trips:
        for i, c in enumerate(trips, start=1):
            _heading(doc, f"Round Trip #{c.get('id', i)}", level=2)
            doc.add_paragraph("Path: " + " -> ".join(str(n) for n in c.get("nodes", [])))
            if c.get("description"):
                doc.add_paragraph(c["description"])
    else:
        doc.add_paragraph("No circular flows detected in this scope.")

    # ------------------------------------------------------------- Money flow
    _heading(doc, "Money Flow")
    defs = report.get("money_flow_definitions") or {}
    if defs.get("overview"):
        doc.add_paragraph(defs["overview"])
    mf = report.get("money_flow", {})
    doc.add_paragraph(
        "Primary destination (where funds accumulate): "
        f"{mf.get('destination_account') or 'None identified'}")

    acc = mf.get("accumulation_accounts", [])[:10]
    if acc:
        _heading(doc, "Accumulation Accounts", level=2)
        if defs.get("accumulation"):
            doc.add_paragraph(defs["accumulation"])
        _table(doc, ["Account", "Total Received", "Senders"],
               [[a.get("node"), _money(a.get("total_received")), a.get("sender_count")]
                for a in acc])

    src = mf.get("source_accounts", [])[:10]
    if src:
        _heading(doc, "Source Accounts", level=2)
        if defs.get("source"):
            doc.add_paragraph(defs["source"])
        _table(doc, ["Account", "Total Sent", "Receivers"],
               [[a.get("node"), _money(a.get("total_sent")), a.get("receiver_count")]
                for a in src])

    lay = mf.get("layering", [])[:10]
    if lay:
        _heading(doc, "Layering / Pass-Through Accounts", level=2)
        if defs.get("layering"):
            doc.add_paragraph(defs["layering"])
        _table(doc, ["Account", "Total In", "Total Out", "Pass-Through"],
               [[a.get("node"), _money(a.get("total_in")), _money(a.get("total_out")),
                 f"{round((a.get('passthrough_ratio') or 0) * 100)}%"] for a in lay])

    # -------------------------------------------- Fund velocity (numbers only)
    timeline = report.get("activity_timeline") or []
    if timeline:
        _heading(doc, "Fund Velocity — Daily Activity")
        _table(doc, ["Date", "Transactions", "Credit", "Debit"],
               [[t.get("date"), t.get("count"), _money(t.get("credit")),
                 _money(t.get("debit"))] for t in timeline])

    # ----------------------------- Cash withdrawal / deposit locations
    cash = report.get("cash_locations") or []
    if cash:
        _heading(doc, "Cash Withdrawal / Deposit Locations")
        doc.add_paragraph(
            "Physical ATM / branch locations where cash was withdrawn (debit) "
            "or deposited (credit), parsed from the transaction narrations — "
            "the on-ground leads for an officer.")
        _table(doc, ["City", "State", "Type", "Amount", "Date", "Time"],
               [[c.get("city"), c.get("state"),
                 "Withdrawal" if c.get("direction") == "DEBIT" else "Deposit",
                 _money(c.get("amount")), c.get("date"), c.get("time")]
                for c in cash])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
