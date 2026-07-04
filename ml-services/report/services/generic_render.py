"""
Generic section-based document renderer — one structured "doc" -> PDF / Excel /
DOCX. Used by the per-service investigation reports (round-trips, money-flow,
money-trail) so each report is laid out consistently and we don't write 3x3
renderers.

Doc shape:
{
  "title": str, "scope": str, "generated_at": str, "subtitle": str (optional),
  "sections": [
     {"heading": str, "kind": "kv",    "data": {k: v}},
     {"heading": str, "kind": "table", "columns": [...], "rows": [[...]],
                       "widths": [..] (optional relative), "right_cols": [i,..]},
     {"heading": str, "kind": "text",  "text": str},
  ]
}
"""

import io

from services.pdf_report import _inr  # Indian-grouped ₹ (ASCII 'Rs')


# --------------------------------------------------------------------------- PDF
def render_pdf(doc) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    title_style = ParagraphStyle("T", parent=styles["Title"], fontSize=17, spaceAfter=4)
    meta_style = ParagraphStyle("M", parent=body, fontSize=9, textColor=colors.HexColor("#52525B"))
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12,
                        textColor=colors.HexColor("#1F4E78"), spaceBefore=6, spaceAfter=4)
    cell = ParagraphStyle("C", parent=body, fontSize=8, leading=10)
    cell_r = ParagraphStyle("CR", parent=cell, alignment=2)
    head_cell = ParagraphStyle("HC", parent=cell, textColor=colors.white, fontName="Helvetica-Bold")

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4, topMargin=16 * mm, bottomMargin=16 * mm,
                            leftMargin=14 * mm, rightMargin=14 * mm)
    avail = pdf.width
    flow = []

    def P(t, s=cell):
        return Paragraph("" if t is None else str(t), s)

    def table(columns, rows, widths=None, right_cols=()):
        ratios = widths or [1] * len(columns)
        total = float(sum(ratios)) or 1.0
        colw = [avail * (r / total) for r in ratios]
        data = [[P(c, head_cell) for c in columns]]
        for row in rows:
            data.append([P(c, cell_r if i in right_cols else cell) for i, c in enumerate(row)])
        t = Table(data, colWidths=colw, repeatRows=1, hAlign="LEFT")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C4D4")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F6FB")]),
        ]))
        return t

    flow.append(Paragraph(doc.get("title", "Investigation Report"), title_style))
    if doc.get("scope"):
        flow.append(Paragraph(doc["scope"], meta_style))
    if doc.get("generated_at"):
        flow.append(Paragraph(f"Generated: {doc['generated_at']}", meta_style))
    if doc.get("subtitle"):
        flow.append(Paragraph(doc["subtitle"], body))
    flow.append(Spacer(1, 8))

    for sec in doc.get("sections", []):
        flow.append(Paragraph(sec.get("heading", ""), h2))
        kind = sec.get("kind")
        if kind == "kv":
            rows = [[k.replace("_", " ").title(), v] for k, v in (sec.get("data") or {}).items()]
            flow.append(table(["Field", "Value"], rows, [2, 3]))
        elif kind == "table":
            rows = sec.get("rows") or []
            if rows:
                flow.append(table(sec.get("columns", []), rows,
                                  sec.get("widths"), tuple(sec.get("right_cols", []))))
            else:
                flow.append(Paragraph(sec.get("empty", "No data."), body))
        elif kind == "text":
            flow.append(Paragraph(sec.get("text", ""), body))
        flow.append(Spacer(1, 8))

    pdf.build(flow)
    return buf.getvalue()


# ------------------------------------------------------------------------- Excel
def render_excel(doc) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(bold=True, color="FFFFFF")
    title_font = Font(bold=True, size=13)

    wb = Workbook()
    wb.remove(wb.active)
    used = set()

    def sheet_name(name):
        # Excel forbids : \ / ? * [ ] in sheet titles, max 31 chars.
        clean = "".join(" " if c in r':\/?*[]' else c for c in (name or "Sheet"))
        base = clean.strip()[:28] or "Sheet"
        n, i = base, 1
        while n in used:
            i += 1
            n = f"{base} {i}"
        used.add(n)
        return n

    info = wb.create_sheet(sheet_name("Overview"))
    info["A1"] = doc.get("title", "Investigation Report")
    info["A1"].font = title_font
    info["A2"] = doc.get("scope", "")
    info["A3"] = f"Generated: {doc.get('generated_at', '')}"
    if doc.get("subtitle"):
        info["A4"] = doc["subtitle"]

    for sec in doc.get("sections", []):
        ws = wb.create_sheet(sheet_name(sec.get("heading", "Section")))
        kind = sec.get("kind")
        if kind == "kv":
            ws.append(["Field", "Value"])
            for c in ws[1]:
                c.font, c.fill = header_font, header_fill
            for k, v in (sec.get("data") or {}).items():
                ws.append([k.replace("_", " ").title(), v])
        elif kind == "table":
            ws.append(sec.get("columns", []))
            for c in ws[1]:
                c.font, c.fill = header_font, header_fill
            for row in sec.get("rows", []):
                ws.append(list(row))
        elif kind == "text":
            ws.append([sec.get("text", "")])
        # autosize
        for col in ws.columns:
            width = max((len(str(c.value)) for c in col if c.value is not None), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(60, width + 2)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# -------------------------------------------------------------------------- DOCX
def render_docx(doc) -> bytes:
    from docx import Document

    d = Document()
    d.add_heading(doc.get("title", "Investigation Report"), level=0)
    if doc.get("scope"):
        d.add_paragraph(doc["scope"])
    if doc.get("generated_at"):
        d.add_paragraph(f"Generated: {doc['generated_at']}")
    if doc.get("subtitle"):
        d.add_paragraph(doc["subtitle"])

    for sec in doc.get("sections", []):
        d.add_heading(sec.get("heading", ""), level=1)
        kind = sec.get("kind")
        if kind == "kv":
            t = d.add_table(rows=0, cols=2)
            t.style = "Light Grid Accent 1"
            for k, v in (sec.get("data") or {}).items():
                cells = t.add_row().cells
                cells[0].text = str(k).replace("_", " ").title()
                cells[1].text = "" if v is None else str(v)
        elif kind == "table":
            cols = sec.get("columns", [])
            rows = sec.get("rows", [])
            if rows and cols:
                t = d.add_table(rows=1, cols=len(cols))
                t.style = "Light Grid Accent 1"
                for i, c in enumerate(cols):
                    t.rows[0].cells[i].text = str(c)
                for row in rows:
                    cells = t.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = "" if val is None else str(val)
            else:
                d.add_paragraph(sec.get("empty", "No data."))
        elif kind == "text":
            d.add_paragraph(sec.get("text", ""))

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def inr(x):
    return _inr(x)
