from docx import Document


# A DOCX bank statement is almost always a table. We only treat a table as the
# transaction table if its header row looks transactional (a date column + a
# money column), so account-summary / address tables aren't turned into rows.
_HEADER_DATE_KEYS = ("date",)
_HEADER_MONEY_KEYS = (
    "amount", "debit", "credit", "balance", "withdrawal", "deposit", "dr", "cr",
)


class DOCXParser:

    def parse(self, file_path: str):
        doc = Document(file_path)

        # 1) Preferred: structured rows from a transaction table (cleanest path,
        #    handled downstream by the same Column Intelligence resolver as
        #    CSV/Excel).
        table_rows = []
        for table in doc.tables:
            table_rows.extend(self._extract_transaction_table(table))

        if table_rows:
            return {
                "source_type": "docx",
                "rows": table_rows,
                "extraction": "tables",
            }

        # 2) Fallback: hand ALL text (paragraphs + every table cell) to the
        #    TextStatementReconstructor as one page-text blob — the same route a
        #    text-based PDF takes, so unusual layouts still yield rows.
        text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                joined = "  ".join(c for c in cells if c)
                if joined:
                    text_lines.append(joined)

        return {
            "source_type": "docx",
            "rows": ["\n".join(text_lines)] if text_lines else [],
            "extraction": "text",
        }

    def _extract_transaction_table(self, table):
        rows = [[cell.text.strip() for cell in r.cells] for r in table.rows]
        if len(rows) < 2:
            return []

        header = rows[0]
        header_joined = " ".join(header).lower()
        has_date = any(k in header_joined for k in _HEADER_DATE_KEYS)
        has_money = any(k in header_joined for k in _HEADER_MONEY_KEYS)
        if not (has_date and has_money):
            return []

        out = []
        for r in rows[1:]:
            if not any(v for v in r):
                continue
            record = {}
            for i, value in enumerate(r):
                key = header[i] if i < len(header) and header[i] else f"col_{i}"
                record[key] = value
            out.append(record)
        return out
